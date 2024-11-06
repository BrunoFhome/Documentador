﻿import os
import tempfile
from io import BytesIO
from flask import Flask, render_template, request, send_file, after_this_request
from docx import Document
from docx.shared import Inches
import requests
import json
from apscheduler.schedulers.background import BackgroundScheduler

app = Flask(__name__)

template_path = os.path.join(os.getcwd(),'doc', 'TemplateDocument.docx')
api_url = os.environ.get('API_URL') or 'http://3.140.207.100/api/getclientes.php'


temp_dir = tempfile.mkdtemp()

scheduler = BackgroundScheduler()
scheduler.add_job(lambda: clean_temp_folder(temp_dir), 'interval', minutes=5)
scheduler.start()

def clean_temp_folder(temp_folder):
    for file_name in os.listdir(temp_folder):
        file_path = os.path.join(temp_folder, file_name)
        try:
            if os.path.isfile(file_path):
                os.unlink(file_path)
        except Exception as e:
            print(f"Error deleting file: {e}")

# Function to fetch data from the API
def get_data_from_api(api_url):
    try:
        response = requests.get(api_url)
        response.raise_for_status()  # Raise HTTPError for bad responses

        # Remove UTF-8 BOM manually
        content = response.content.decode('utf-8')
        if content.startswith('\ufeff'):
            content = content[1:]

        data = json.loads(content)

        return data
    except requests.RequestException as e:
        print(f"Error in API request: {e}")
        return None

@app.route('/')
def index():
    # Fetch data from the API
    clients_data = get_data_from_api(api_url)

    # Print the received data for debugging
    #print("Received clients data:", clients_data)

    # Pass the data to the HTML template
    return render_template('index.html', clients=clients_data)

@app.route('/process_template', methods=['POST'])
def process_template():
    doc = Document(template_path)

    # Replace placeholders with data from the form
    replace_placeholder(doc, '@chamado', request.form['data1'])
    replace_placeholder(doc, '@cliente', request.form['data2'])
    replace_placeholder(doc, '@modulo', request.form['data3'])
    replace_placeholder(doc, '@data', request.form['data4'])
    replace_placeholder(doc, '@descricao', request.form['data5'])

    # Save the modified document
    modified_filename = f'DOCUMENTAÇÃO - {request.form["data2"]}.docx'
    modified_path = os.path.join(temp_dir, modified_filename)
    doc.save(modified_path)

    # Handle multiple image uploads and their descriptions
    images = request.files.getlist('images[]')
    for index, image in enumerate(images):
        description = request.form.get(f'image_description_{index}', '')
        image_filename = f'image_{index}_{image.filename}'
        image_path = os.path.join(temp_dir, image_filename)
        image.save(image_path)
        # Add image and description to the document
        doc.add_paragraph(description)
        doc.add_picture(image_path)

    # Save the final document with images
    doc.save(modified_path)

    # return 'Template processed successfully'
    #return send_file(modified_path, as_attachment=True, attachment_filename=modified_filename)
    return send_file(
        modified_path,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=modified_filename
    )
    

def replace_placeholder(doc, placeholder, replacement):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, replacement)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, replacement)

def save_image(file):
    image_folder = os.path.join(temp_dir)
    if not os.path.exists(image_folder):
        os.makedirs(image_folder)

    image_path = os.path.join(image_folder, file.filename)
    file.save(image_path)
    return image_path

def insert_images_after_placeholder(doc_path, image_paths, placeholder='@prints'):
    doc = Document(doc_path)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder in paragraph.text:
                        # Clear the existing paragraph with the @prints placeholder
                        paragraph.clear()

                        # Insert new paragraphs with images after the cleared one
                        for image_path in image_paths:
                            p = cell.add_paragraph()
                            run = p.add_run()
                            run.add_picture(image_path, width=Inches(1.0))

    # Save the modified document
    doc.save(doc_path)

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
